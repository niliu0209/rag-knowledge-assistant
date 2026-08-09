"""测试合成文档生成（本地隔离、可复现，不依赖真实用户文件）。

- 中文 docx：python-docx 直接写 XML，无需系统字体
- 英文 PDF：reportlab（有文本层；解析管线行为与中文一致，中文质量手动验收）
- 图片型 PDF：pypdf 空白页（无文本层，模拟扫描件）
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas


def make_docx(path: Path, text: str | None = None) -> Path:
    """生成含中文段落的小 docx。"""
    doc = Document()
    doc.add_paragraph(
        text or "行政部周工作小结：本周完成办公公共区域卫生巡检三次，协调保洁补齐洗手液等基础耗材；"
        "完成 2 名新入职员工工位物资配齐；维护会议室预约系统，协调化解时段预约冲突 2 起。"
    )
    doc.add_paragraph("下周计划：启动 8 月全品类办公物资月度盘点；简化修订工位物资线上申领流程。")
    doc.save(path)
    return path


def make_pdf(path: Path) -> Path:
    """生成 2 页英文文本 PDF（有文本层，含表格字样供核对）。"""
    c = canvas.Canvas(str(path))
    c.drawString(72, 750, "Office supplies procurement summary July 2026.")
    c.drawString(72, 730, "Table: A4 paper 500 sheets per pack, price 25 yuan.")
    c.showPage()
    c.drawString(72, 750, "Second page: black gel pens 60 units, 2 yuan each.")
    c.save()
    return path


def make_image_pdf(path: Path) -> Path:
    """生成无文本层 PDF（空白页，等价于扫描件/图片型）。"""
    import pypdf

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with open(path, "wb") as f:
        writer.write(f)
    return path
